"""
notes_to_docx — раскладывает фото-конспекты по 2 штуки на лист A4 (landscape)
в формате docx с порядком страниц как в тетради (booklet imposition).

Два режима:
  1) Фото уже названы по порядку (page1_print.jpeg ...) — вставляются как есть.
  2) Произвольные фото — Gemini распознаёт текст, определяет порядок,
     текст рендерится в изображение через пользовательский TTF-шрифт (Pillow),
     и уже эти отрендеренные картинки вставляются в docx. Так шрифт
     гарантированно отображается одинаково на любой машине.

Логика раскладки (booklet imposition):
  N страниц -> добивается до кратного 4 «пустыми» местами,
  пары формируются так, чтобы при складывании листа пополам
  страницы шли по порядку.
"""

from __future__ import annotations

import os
import re
import sys
import tempfile
from dataclasses import dataclass
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor, as_completed

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Cm
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


IMAGE_EXTS = (".jpg", ".jpeg", ".png", ".bmp", ".tif", ".tiff", ".webp")
DEFAULT_WIDTH_CM = 14.4
DEFAULT_HEIGHT_CM = 18.41
DEFAULT_FONT_SIZE_PT = 14
DEFAULT_DPI = 200
DEFAULT_LINE_SPACING = 1.35
DEFAULT_MARGIN_CM = 0.3


# ---------------------------------------------------------------------------
# Логирование
# ---------------------------------------------------------------------------

def log(msg: str = "") -> None:
    print(msg, flush=True)


def log_inline(msg: str) -> None:
    print(msg, end="", flush=True)


# ---------------------------------------------------------------------------
# Утилиты
# ---------------------------------------------------------------------------

def natural_key(name: str):
    return [int(p) if p.isdigit() else p.lower() for p in re.split(r"(\d+)", name)]


def list_images(directory: Path) -> list[Path]:
    files = [p for p in directory.iterdir()
             if p.is_file() and p.suffix.lower() in IMAGE_EXTS]
    files.sort(key=lambda p: natural_key(p.name))
    return files


def compute_dims(image_path: Path,
                 width_cm: float | None,
                 height_cm: float | None) -> tuple[float, float]:
    if width_cm and height_cm:
        return width_cm, height_cm
    with Image.open(image_path) as img:
        iw, ih = img.size
    aspect = iw / ih
    if width_cm and not height_cm:
        return width_cm, width_cm / aspect
    if height_cm and not width_cm:
        return height_cm * aspect, height_cm
    return DEFAULT_WIDTH_CM, DEFAULT_WIDTH_CM / aspect


# ---------------------------------------------------------------------------
# Режимы разбивки на листы
# ---------------------------------------------------------------------------

LAYOUT_BOOKLET    = "booklet"
LAYOUT_SEQUENTIAL = "sequential"
LAYOUT_SINGLE     = "single"


def layout_booklet(num_pages: int) -> list[tuple[int | None, int | None]]:
    """Буклет: листы складываются стопкой и сшиваются посередине.

    Лист 1 лицо:  [стр.1] | [стр.4]
    Лист 1 оборот:[стр.2]  | [стр.3]
    ...

    Порядок чтения корректен при складывании. Остаток (< 4 стр.) идёт
    последовательно без пустых страниц.

    Пример 4 стр: [(1,4),(2,3)]
    Пример 7 стр: [(1,4),(2,3),(5,None),(6,7)]
    Пример 8 стр: [(1,8),(2,7),(3,6),(4,5)]
    Пример 9 стр: [(1,8),(2,7),(3,6),(4,5),(9,None)]
    """
    if num_pages <= 0:
        return []
    full_groups = num_pages // 4
    remainder = num_pages % 4
    sheets: list[tuple[int | None, int | None]] = []

    # Буклетная раскладка для полных групп по 4 страницы
    for g in range(full_groups):
        base = g * 4
        lo = base + 1
        hi = base + 4
        sheets.append((lo, hi))
        sheets.append((lo + 1, hi - 1))

    # Остаток: последовательная раскладка без пустых страниц
    if remainder:
        start = full_groups * 4 + 1
        for i in range(start, num_pages + 1, 2):
            left = i
            right = i + 1 if i + 1 <= num_pages else None
            sheets.append((left, right))

    return sheets


def layout_sequential(num_pages: int) -> list[tuple[int | None, int | None]]:
    """Последовательный разворот: стр.1+2, стр.3+4, стр.5+6 ...

    Печатается двусторонне: лицо листа = левая страница разворота,
    оборот = правая. Потом листы брошюруются по краю (скоросшиватель,
    пружина, скрепка сбоку).

    Работает с любым числом страниц. Пустое место — только на
    последнем листе если страниц нечётное число.

    Пример 9 стр: [(1,2),(3,4),(5,6),(7,8),(9,None)]
    """
    if num_pages <= 0:
        return []
    sheets: list[tuple[int | None, int | None]] = []
    for i in range(0, num_pages, 2):
        left = i + 1
        right = i + 2 if i + 2 <= num_pages else None
        sheets.append((left, right))
    return sheets


def layout_single(num_pages: int) -> list[tuple[int | None, int | None]]:
    """Одна страница по центру листа A4.

    Правая ячейка таблицы остаётся пустой. Подходит для просмотра
    на экране или печати с увеличенным масштабом.

    Пример 4 стр: [(None,1),(None,2),(None,3),(None,4)]
    """
    if num_pages <= 0:
        return []
    return [(None, i + 1) for i in range(num_pages)]


def make_pairs(num_pages: int,
               layout: str) -> list[tuple[int | None, int | None]]:
    if layout == LAYOUT_SEQUENTIAL:
        return layout_sequential(num_pages)
    if layout == LAYOUT_SINGLE:
        return layout_single(num_pages)
    return layout_booklet(num_pages)  # default


# backward-compat alias
def booklet_pairs(num_pages: int) -> list[tuple[int | None, int | None]]:
    return layout_booklet(num_pages)


# ---------------------------------------------------------------------------
# Рендер текста в изображение (замена текста-в-docx)
# ---------------------------------------------------------------------------

def _measure_rendered_width(text: str, font: ImageFont.FreeTypeFont) -> int:
    """Измеряет реальную ширину отрендеренного текста через пиксели."""
    if not text:
        return 0
    test_img = Image.new("L", (2000, 100), 255)
    test_draw = ImageDraw.Draw(test_img)
    test_draw.text((0, 0), text, fill=0, font=font)
    pixels = test_img.load()
    w, h = test_img.size
    right = 0
    for x in range(w - 1, -1, -1):
        for y in range(h):
            if pixels[x, y] != 255:
                right = x
                break
        if right:
            break
    return right + 1 if right else 0


def _measure_space(font: ImageFont.FreeTypeFont, draw: ImageDraw.ImageDraw) -> int:
    """Возвращает реальную ширину пробела.

    Сначала пробует textlength (работает для шрифтов с хорошими метриками).
    Если метрики сломаны (все widths = 0), измеряет через пиксели
    и вычисляет пробел как долю от средней ширины символа.
    """
    w_ab = draw.textlength("аб", font=font)
    w_asb = draw.textlength("а б", font=font)
    space_w = w_asb - w_ab
    if space_w >= 2:
        return int(space_w)

    sample = "абвгдежзиклмнопрстуфхцчшщэюя"
    total = 0
    count = 0
    for ch in sample:
        cw = _measure_rendered_width(ch, font)
        if cw > 0:
            total += cw
            count += 1
    avg = total / count if count else 10
    return max(6, int(avg * 0.5))


def _text_width(text: str, font: ImageFont.FreeTypeFont,
                draw: ImageDraw.ImageDraw, space_w: int) -> int:
    """Ширина строки с учётом реального размера пробела."""
    if " " not in text:
        w = draw.textlength(text, font=font)
        if w > 0:
            return int(w)
        return _measure_rendered_width(text, font)
    parts = text.split(" ")
    total = 0
    for p in parts:
        w = draw.textlength(p, font=font)
        if w > 0:
            total += int(w)
        else:
            total += _measure_rendered_width(p, font)
    total += space_w * (len(parts) - 1)
    return total


def _word_wrap(text: str, font: ImageFont.FreeTypeFont,
               max_width: int, draw: ImageDraw.ImageDraw,
               space_w: int) -> list[str]:
    """Разбивает текст на строки, укладывающиеся в max_width пикселей."""
    result: list[str] = []
    for paragraph in text.split("\n"):
        if not paragraph.strip():
            result.append("")
            continue
        words = paragraph.split()
        if not words:
            result.append("")
            continue
        current = words[0]
        for word in words[1:]:
            test = current + " " + word
            if _text_width(test, font, draw, space_w) <= max_width:
                current = test
            else:
                result.append(current)
                current = word
        result.append(current)
    return result


def render_text_to_image(
    text: str,
    ttf_path: Path,
    width_cm: float = DEFAULT_WIDTH_CM,
    height_cm: float = DEFAULT_HEIGHT_CM,
    font_size_pt: float = DEFAULT_FONT_SIZE_PT,
    dpi: int = DEFAULT_DPI,
    margin_cm: float = DEFAULT_MARGIN_CM,
    line_spacing: float = DEFAULT_LINE_SPACING,
) -> Image.Image:
    """Рендерит текст в изображение белого фона используя TTF-шрифт.

    Результат — PNG-картинка, которая затем вставляется в docx как обычное фото.
    """
    px_per_cm = dpi / 2.54
    img_w = int(width_cm * px_per_cm)
    img_h = int(height_cm * px_per_cm)
    margin_px = int(margin_cm * px_per_cm)

    font_size_px = int(font_size_pt * dpi / 72)
    font = ImageFont.truetype(str(ttf_path), size=font_size_px)

    img = Image.new("RGB", (img_w, img_h), "white")
    draw = ImageDraw.Draw(img)

    space_w = _measure_space(font, draw)
    text_area_w = img_w - 2 * margin_px
    lines = _word_wrap(text or "", font, text_area_w, draw, space_w)

    step = int(font_size_px * line_spacing)
    y = margin_px
    bottom = img_h - margin_px

    for line in lines:
        if y + step > bottom:
            break
        # Рендерим слово за словом с явным отступом space_w между ними.
        # Это обходит баг шрифтов с нулевым пробелом.
        x = margin_px
        words = line.split(" ")
        for wi, word in enumerate(words):
            if word:
                draw.text((x, y), word, fill="black", font=font)
                ww = draw.textlength(word, font=font)
                if ww > 0:
                    x += int(ww)
                else:
                    x += _measure_rendered_width(word, font)
            if wi < len(words) - 1:
                x += space_w
        y += step

    return img


def render_texts_to_images(
    texts: list[str],
    ttf_path: Path,
    out_dir: Path,
    width_cm: float,
    height_cm: float,
    font_size_pt: float,
) -> list[Path]:
    """Рендерит список текстов в PNG-файлы. Возвращает пути."""
    rendered_dir = out_dir / "rendered"
    rendered_dir.mkdir(exist_ok=True)
    paths: list[Path] = []
    for i, text in enumerate(texts, 1):
        img = render_text_to_image(
            text, ttf_path,
            width_cm=width_cm,
            height_cm=height_cm,
            font_size_pt=font_size_pt,
        )
        p = rendered_dir / f"page_{i:03d}.png"
        img.save(str(p), "PNG")
        paths.append(p)
        log(f"  [{i}/{len(texts)}] -> {p.name}")
    return paths


def split_text_to_pages(
    text: str,
    font_path: Path,
    width_cm: float = DEFAULT_WIDTH_CM,
    height_cm: float = DEFAULT_HEIGHT_CM,
    font_size_pt: float = DEFAULT_FONT_SIZE_PT,
    dpi: int = DEFAULT_DPI,
    margin_cm: float = DEFAULT_MARGIN_CM,
    line_spacing: float = DEFAULT_LINE_SPACING,
) -> list[str]:
    """Разбивает текст на страницы по размеру A4 с учётом шрифта."""
    px_per_cm = dpi / 2.54
    img_w = int(width_cm * px_per_cm)
    img_h = int(height_cm * px_per_cm)
    margin_px = int(margin_cm * px_per_cm)

    font_size_px = int(font_size_pt * dpi / 72)
    font = ImageFont.truetype(str(font_path), size=font_size_px)
    draw = ImageDraw.Draw(Image.new("RGB", (img_w, img_h)))

    space_w = _measure_space(font, draw)
    text_area_w = img_w - 2 * margin_px
    step = int(font_size_px * line_spacing)
    text_area_h = img_h - 2 * margin_px
    max_lines = max(1, text_area_h // step)

    lines = _word_wrap(text, font, text_area_w, draw, space_w)

    pages: list[str] = []
    for i in range(0, len(lines), max_lines):
        chunk = lines[i:i + max_lines]
        pages.append("\n".join(chunk))

    return pages if pages else [""]


# ---------------------------------------------------------------------------
# Markdown поддержка
# ---------------------------------------------------------------------------

_HEADING_RE = re.compile(r"^(#{1,3})\s+(.*)")
_CENTER_RE = re.compile(r"^>\s+(.*)")
_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_IGNORE_RE = re.compile(r"^---$")


@dataclass
class MdBlock:
    kind: str  # "heading" | "center" | "text" | "blank"
    text: str = ""
    level: int = 1


def parse_markdown(text: str) -> list[MdBlock]:
    """Парсит markdown в список блоков.

    Поддерживаемый синтаксис:
      # Заголовок      — крупный шрифт, по центру
      ## Подзаголовок  — средний шрифт, по центру
      ### Мелкий       — чуть крупнее обычного, по центру
      > текст          — по центру
      ---              — пустая строка
      **жирный**       — убирает маркеры, текст остаётся
    """
    blocks: list[MdBlock] = []
    for raw_line in text.split("\n"):
        line = raw_line.rstrip()

        if _IGNORE_RE.match(line):
            blocks.append(MdBlock(kind="blank"))
            continue

        m = _HEADING_RE.match(line)
        if m:
            level = len(m.group(1))
            t = _BOLD_RE.sub(r"\1", m.group(2)).strip()
            blocks.append(MdBlock(kind="heading", text=t, level=level))
            continue

        m = _CENTER_RE.match(line)
        if m:
            t = _BOLD_RE.sub(r"\1", m.group(1)).strip()
            blocks.append(MdBlock(kind="center", text=t))
            continue

        t = _BOLD_RE.sub(r"\1", line).strip()
        if not t:
            blocks.append(MdBlock(kind="blank"))
        else:
            blocks.append(MdBlock(kind="text", text=t))

    return blocks


def _heading_scale(level: int) -> float:
    """Множитель размера шрифта для заголовка."""
    return {1: 1.6, 2: 1.3, 3: 1.1}.get(level, 1.0)


def render_markdown_to_image(
    blocks: list[MdBlock],
    ttf_path: Path,
    base_size_pt: float = DEFAULT_FONT_SIZE_PT,
    width_cm: float = DEFAULT_WIDTH_CM,
    height_cm: float = DEFAULT_HEIGHT_CM,
    dpi: int = DEFAULT_DPI,
    margin_cm: float = DEFAULT_MARGIN_CM,
    line_spacing: float = DEFAULT_LINE_SPACING,
) -> Image.Image:
    """Рендерит распарсенные markdown-блоки в изображение."""
    px_per_cm = dpi / 2.54
    img_w = int(width_cm * px_per_cm)
    img_h = int(height_cm * px_per_cm)
    margin_px = int(margin_cm * px_per_cm)

    img = Image.new("RGB", (img_w, img_h), "white")
    draw = ImageDraw.Draw(img)
    text_area_w = img_w - 2 * margin_px
    y = margin_px
    bottom = img_h - margin_px

    for block in blocks:
        if block.kind == "blank":
            y += int(base_size_pt * dpi / 72 * line_spacing * 0.5)
            if y > bottom:
                break
            continue

        scale = _heading_scale(block.level) if block.kind == "heading" else 1.0
        font_size_px = int(base_size_pt * scale * dpi / 72)
        font = ImageFont.truetype(str(ttf_path), size=font_size_px)
        step = int(font_size_px * line_spacing)

        if y + step > bottom:
            break

        space_w = _measure_space(font, draw)
        align = "center" if block.kind in ("heading", "center") else "left"
        lines = _word_wrap(block.text, font, text_area_w, draw, space_w)

        for line in lines:
            if y + step > bottom:
                break
            if align == "center":
                lw = _text_width(line, font, draw, space_w)
                x = margin_px + max(0, (text_area_w - lw) // 2)
            else:
                x = margin_px
            words = line.split(" ")
            for wi, word in enumerate(words):
                if word:
                    draw.text((x, y), word, fill="black", font=font)
                    ww = draw.textlength(word, font=font)
                    if ww > 0:
                        x += int(ww)
                    else:
                        x += _measure_rendered_width(word, font)
                if wi < len(words) - 1:
                    x += space_w
            y += step

    return img


def split_markdown_to_pages(
    text: str,
    ttf_path: Path,
    base_size_pt: float = DEFAULT_FONT_SIZE_PT,
    width_cm: float = DEFAULT_WIDTH_CM,
    height_cm: float = DEFAULT_HEIGHT_CM,
    dpi: int = DEFAULT_DPI,
    margin_cm: float = DEFAULT_MARGIN_CM,
    line_spacing: float = DEFAULT_LINE_SPACING,
) -> list[list[MdBlock]]:
    """Разбивает markdown на страницы (список списков блоков)."""
    blocks = parse_markdown(text)
    px_per_cm = dpi / 2.54
    img_w = int(width_cm * px_per_cm)
    img_h = int(height_cm * px_per_cm)
    margin_px = int(margin_cm * px_per_cm)
    text_area_w = img_w - 2 * margin_px
    text_area_h = img_h - 2 * margin_px

    pages: list[list[MdBlock]] = []
    current_page: list[MdBlock] = []
    y = 0

    for block in blocks:
        if block.kind == "blank":
            est = int(base_size_pt * dpi / 72 * line_spacing * 0.5)
            if y + est > text_area_h and current_page:
                pages.append(current_page)
                current_page = []
                y = 0
            y += est
            current_page.append(block)
            continue

        scale = _heading_scale(block.level) if block.kind == "heading" else 1.0
        font_size_px = int(base_size_pt * scale * dpi / 72)
        font = ImageFont.truetype(str(ttf_path), size=font_size_px)
        step = int(font_size_px * line_spacing)
        space_w = _measure_space(font, ImageDraw.Draw(Image.new("RGB", (1, 1))))
        lines = _word_wrap(block.text, font, text_area_w, ImageDraw.Draw(Image.new("RGB", (1, 1))), space_w)

        needed = len(lines) * step
        if y + needed > text_area_h and current_page:
            pages.append(current_page)
            current_page = []
            y = 0

        current_page.append(block)
        y += needed

    if current_page:
        pages.append(current_page)

    return pages if pages else [[]]


# ---------------------------------------------------------------------------
# Генерация docx с изображениями (общая для обоих режимов)
# ---------------------------------------------------------------------------

def setup_landscape_a4(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(0.3)
    section.right_margin = Cm(0.3)
    section.top_margin = Cm(0.5)
    section.bottom_margin = Cm(0.5)


def _zero_cell_margins(cell) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn("w:tcMar"))
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for side in ("top", "left", "bottom", "right"):
        node = tcMar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tcMar.append(node)
        node.set(qn("w:w"), "0")
        node.set(qn("w:type"), "dxa")


def _set_fixed_table_layout(table) -> None:
    tblPr = table._element.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        table._element.insert(0, tblPr)
    layout = tblPr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")


def create_sheet_docx(left_img: Path | None,
                      right_img: Path | None,
                      output_path: Path,
                      width_cm: float | None,
                      height_cm: float | None) -> None:
    doc = Document()
    setup_landscape_a4(doc)

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    _set_fixed_table_layout(table)

    half_page = Cm(14.55)

    for cell, img_path in zip(table.rows[0].cells, [left_img, right_img]):
        cell.width = half_page
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _zero_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.left_indent = Cm(0)
        pf.right_indent = Cm(0)
        pf.space_before = Cm(0)
        pf.space_after = Cm(0)
        run = p.add_run()
        if img_path is not None:
            w, h = compute_dims(img_path, width_cm, height_cm)
            run.add_picture(str(img_path), width=Cm(w), height=Cm(h))

    doc.save(str(output_path))


# ---------------------------------------------------------------------------
# Интерактивный ввод
# ---------------------------------------------------------------------------

def prompt_float(label: str, default: float | None) -> float | None:
    default_str = "" if default is None else f"{default}"
    suffix = f" [{default_str}]" if default_str else " [auto]"
    raw = input(f"{label}{suffix}: ").strip().replace(",", ".")
    if raw == "":
        return default
    if raw == "-":
        return None
    try:
        return float(raw)
    except ValueError:
        log(f"  не понял '{raw}', использую {default}")
        return default


def ask_dimensions() -> tuple[float | None, float | None]:
    log()
    log("Размер фото (соотношение сторон сохраняется).")
    log("Введите ОДНУ сторону - вторая посчитается автоматически.")
    log("Чтобы оставить сторону 'авто' - введите '-'.")
    log(f"По умолчанию: ширина {DEFAULT_WIDTH_CM} см, высота {DEFAULT_HEIGHT_CM} см.")
    width = prompt_float("Ширина (см)", DEFAULT_WIDTH_CM)
    height = prompt_float("Высота (см)", DEFAULT_HEIGHT_CM)
    return width, height


def ask_render_params() -> tuple[float, float, float]:
    """Размеры отрендеренной страницы и кегль шрифта для режима 2."""
    log()
    log("Размер страницы и кегль шрифта для рендера.")
    log(f"По умолчанию: {DEFAULT_WIDTH_CM} x {DEFAULT_HEIGHT_CM} см, {DEFAULT_FONT_SIZE_PT} pt.")
    width = prompt_float("Ширина страницы (см)", DEFAULT_WIDTH_CM)
    height = prompt_float("Высота страницы (см)", DEFAULT_HEIGHT_CM)
    size = prompt_float("Кегль шрифта (pt)", float(DEFAULT_FONT_SIZE_PT))
    return (
        width or DEFAULT_WIDTH_CM,
        height or DEFAULT_HEIGHT_CM,
        size or float(DEFAULT_FONT_SIZE_PT),
    )


def resolve_font_path() -> Path:
    """Получает путь к TTF: сначала из .env, потом спрашивает интерактивно."""
    env_font = ""
    try:
        import ai_client
        env_font = ai_client.load_config().get("font_path", "")
    except Exception:
        pass

    if env_font:
        p = Path(os.path.expanduser(env_font)).resolve()
        if p.is_file() and p.suffix.lower() in (".ttf", ".otf"):
            log(f"TTF-шрифт (из .env): {p}")
            return p
        log(f"  FONT_PATH из .env не найден или не TTF: {env_font}")

    log()
    log("Путь к TTF-файлу вашего рукописного шрифта.")
    if env_font:
        log(f"  (в .env указано: {env_font})")
    while True:
        raw = input("Путь к .ttf: ").strip()
        if not raw and env_font:
            raw = env_font
        if not raw:
            log("  путь обязателен")
            continue
        path = Path(os.path.expanduser(raw)).resolve()
        if not path.is_file():
            log(f"  файл не найден: {path}")
            continue
        if path.suffix.lower() not in (".ttf", ".otf"):
            log(f"  ожидался .ttf или .otf, получено: {path.suffix}")
            continue
        return path


def ask_directory() -> Path | None:
    raw = input("Директория с фото: ").strip()
    if not raw:
        log("Не указана директория.")
        return None
    directory = Path(os.path.expanduser(raw)).resolve()
    if not directory.is_dir():
        log(f"Не нашёл директорию: {directory}")
        return None
    return directory


def ask_layout() -> str:
    """Спрашивает режим разбивки на листы."""
    log()
    log("Тип печати:")
    log("  1) Буклет      - складываешь листы стопкой, сшиваешь посередине")
    log("                   (нужно кратное 4 стр., пустые добавляются авто)")
    log("  2) Разворот    - стр.1+2, 3+4, 5+6 ... печать двусторонняя,")
    log("                   брошюруешь по краю (скоросшиватель/пружина)")
    log("                   (работает с любым кол-вом стр., пустое только в конце)")
    log("  3) По одной    - каждая страница отдельно по центру A4")
    while True:
        raw = input("Тип печати [1/2/3] (по умолчанию 1): ").strip()
        if raw in ("", "1"):
            return LAYOUT_BOOKLET
        if raw == "2":
            return LAYOUT_SEQUENTIAL
        if raw == "3":
            return LAYOUT_SINGLE
        log("  введите 1, 2 или 3")


def ask_mode() -> int:
    backend_info = ""
    try:
        import ai_client
        cfg = ai_client.load_config()
        backend_info = f" [{ai_client.backend_label(cfg)}]"
    except Exception:
        backend_info = " [настройте .env]"
    log("Режимы работы:")
    log("  1) Фото уже названы по порядку - в docx идут сами фото")
    log(f"  2) Произвольные фото - OCR + рендер текста вашим TTF в картинки -> docx{backend_info}")
    log("  3) Текстовый файл print.txt - рендер текста вашим TTF в картинки -> docx")
    while True:
        raw = input("Выберите режим [1/2/3] (по умолчанию 1): ").strip()
        if raw == "" or raw == "1":
            return 1
        if raw == "2":
            return 2
        if raw == "3":
            return 3
        log("  введите 1, 2 или 3")


# ---------------------------------------------------------------------------
# Gemini pipeline
# ---------------------------------------------------------------------------

def _preview(text: str, limit: int = 60) -> str:
    one_line = re.sub(r"\s+", " ", (text or "").strip())
    if len(one_line) <= limit:
        return one_line
    return one_line[:limit] + "..."


def ai_pipeline(images: list[Path], out_dir: Path) -> tuple[list[Path], list[str]]:
    """OCR каждого фото через выбранный бэкенд + определение порядка страниц.

    Возвращает (упорядоченные пути к оригиналам, упорядоченные тексты).
    """
    import ai_client

    cfg = ai_client.load_config()
    log()
    log(f"Использую: {ai_client.backend_label(cfg)}")
    if cfg["proxy"]:
        log(f"Прокси: {cfg['proxy']}")

    # Загружаем кэш
    cache = ai_client.load_cache(out_dir)
    
    log()
    log("[1/2] Распознаю текст (параллельно)...")
    
    results: dict[Path, str] = {}
    to_process = []
    
    for img in images:
        if img.name in cache:
            results[img] = cache[img.name]
        else:
            to_process.append(img)

    if to_process:
        log(f"  Нужно распознать: {len(to_process)} из {len(images)}")
        
        # Определяем кол-во воркеров (для Gemini обычно не больше 2-5 в бесплатке)
        max_workers = 3 if cfg["backend"] == "gemini" else 5
        
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            future_to_img = {
                executor.submit(ai_client.recognize_notes, img, cfg=cfg): img 
                for img in to_process
            }
            
            completed = 0
            for future in as_completed(future_to_img):
                img = future_to_img[future]
                completed += 1
                try:
                    text = future.result()
                    results[img] = text
                    cache[img.name] = text
                    # Сохраняем кэш после каждого успешного распознавания
                    ai_client.save_cache(out_dir, cache)
                    log(f"  [{completed}/{len(to_process)}] {img.name} OK: \"{_preview(text)}\"")
                except Exception as exc:
                    log(f"  [{completed}/{len(to_process)}] {img.name} ОШИБКА: {exc}")
                    # В случае фатальной ошибки можно прервать, но лучше собрать что можно
    else:
        log("  Все страницы найдены в кэше.")

    # Собираем тексты в исходном порядке для передачи в order_pages
    texts_in_input_order = [results.get(img, "") for img in images]

    log()
    log("[2/2] Определяю порядок страниц...")
    try:
        order = ai_client.order_pages(texts_in_input_order, cfg=cfg)
    except Exception as exc:
        log(f"  [!] Не удалось определить порядок: {exc}")
        log("      Использую исходный порядок.")
        order = list(range(len(images)))

    log(f"  Порядок (индексы): {order}")
    ordered_imgs = [images[i] for i in order]
    ordered_texts = [texts_in_input_order[i] for i in order]

    log()
    log("Итоговый порядок файлов:")
    for k, img in enumerate(ordered_imgs, 1):
        log(f"  {k:>3}. {img.name}")
    return ordered_imgs, ordered_texts


# ---------------------------------------------------------------------------
# Генерация booklet docx (общая для обоих режимов)
# ---------------------------------------------------------------------------

def generate_booklet(page_images: list[Path],
                     out_dir: Path,
                     width_cm: float | None,
                     height_cm: float | None,
                     layout: str = LAYOUT_BOOKLET) -> list[Path]:
    """Создаёт docx из списка изображений-страниц. Возвращает пути к docx."""
    sheets = make_pairs(len(page_images), layout)
    log()
    log(f"Получится файлов: {len(sheets)}")
    log()
    created: list[Path] = []
    for idx, (left_idx, right_idx) in enumerate(sheets, start=1):
        left_img = page_images[left_idx - 1] if left_idx else None
        right_img = page_images[right_idx - 1] if right_idx else None
        left_label = f"p{left_idx}" if left_idx else "blank"
        right_label = f"p{right_idx}" if right_idx else "blank"
        out_path = out_dir / f"sheet_{idx:02d}_{left_label}_{right_label}.docx"
        create_sheet_docx(left_img, right_img, out_path, width_cm, height_cm)
        log(f"  -> {out_path.name}")
        created.append(out_path)
    return created


# ---------------------------------------------------------------------------
# main
# ---------------------------------------------------------------------------

def main() -> int:
    log("notes_to_docx - фото-конспекты по 2 на лист A4")
    log()

    mode = ask_mode()
    layout = ask_layout()
    log()

    directory = ask_directory()
    if directory is None:
        return 1

    out_dir = directory / "docx_output"
    out_dir.mkdir(exist_ok=True)

    # ==================== MODE 1 ====================
    if mode == 1:
        images = list_images(directory)
        if not images:
            log("В директории нет изображений.")
            return 1

        log()
        log(f"Найдено фото: {len(images)}")
        for i, p in enumerate(images, 1):
            log(f"  {i:>3}. {p.name}")

        width_cm, height_cm = ask_dimensions()
        generate_booklet(images, out_dir, width_cm, height_cm, layout=layout)
        log()
        log(f"Готово. Файлы здесь: {out_dir}")
        return 0

    # ==================== MODE 3 ====================
    if mode == 3:
        # 1. Шрифт
        try:
            font_path = resolve_font_path()
            font = ImageFont.truetype(str(font_path), size=32)
            family, _ = font.getname()
            log(f"Шрифт: {family!r} ({font_path.name})")
        except Exception as exc:
            log(f"\nНе смог загрузить шрифт: {exc}")
            return 3

        # 2. Чтение текстового файла (print.md или print.txt)
        print_md = directory / "print.md"
        print_txt = directory / "print.txt"
        use_markdown = False

        if print_md.is_file():
            text = print_md.read_text(encoding="utf-8").strip()
            use_markdown = True
            log(f"Прочитано {len(text)} символов из print.md (markdown-режим)")
        elif print_txt.is_file():
            text = print_txt.read_text(encoding="utf-8").strip()
            log(f"Прочитано {len(text)} символов из print.txt")
        else:
            log(f"\nФайлы print.md и print.txt не найдены в {directory}.")
            log("Создайте print.md (с заголовками и центровкой) или print.txt (простой текст).")
            return 4

        if not text:
            log("\nФайл пуст.")
            return 4

        # 3. Параметры рендера
        width_cm, height_cm, font_size_pt = ask_render_params()

        # 4. Разбивка текста на страницы
        log()
        rendered_dir = out_dir / "rendered"
        rendered_dir.mkdir(exist_ok=True)
        rendered_images: list[Path] = []

        if use_markdown:
            log("Разбиваю markdown на страницы...")
            md_pages = split_markdown_to_pages(
                text, font_path,
                base_size_pt=font_size_pt,
                width_cm=width_cm,
                height_cm=height_cm,
            )
            log(f"Получится {len(md_pages)} стр.")
            log()
            log("Рендерю markdown в изображения через ваш шрифт...")
            for i, blocks in enumerate(md_pages, 1):
                img = render_markdown_to_image(
                    blocks, font_path,
                    base_size_pt=font_size_pt,
                    width_cm=width_cm,
                    height_cm=height_cm,
                )
                p = rendered_dir / f"page_{i:03d}.png"
                img.save(str(p), "PNG")
                rendered_images.append(p)
                log(f"  [{i}/{len(md_pages)}] -> {p.name}")
        else:
            log("Разбиваю текст на страницы...")
            pages = split_text_to_pages(
                text, font_path,
                width_cm=width_cm,
                height_cm=height_cm,
                font_size_pt=font_size_pt,
            )
            log(f"Получится {len(pages)} стр.")
            log()
            log("Рендерю текст в изображения через ваш шрифт...")
            rendered_images = render_texts_to_images(
                pages, font_path, out_dir,
                width_cm=width_cm,
                height_cm=height_cm,
                font_size_pt=font_size_pt,
            )

        # 5. Booklet docx из отрендеренных картинок
        generate_booklet(rendered_images, out_dir, width_cm, height_cm, layout=layout)

        log()
        log(f"Готово. Файлы здесь: {out_dir}")
        log(f"Отрендеренные страницы: {out_dir / 'rendered'}")
        return 0

    # ==================== MODE 2 ====================
    images = list_images(directory)
    if not images:
        log("В директории нет изображений.")
        return 1

    log()
    log(f"Найдено фото: {len(images)}")
    for i, p in enumerate(images, 1):
        log(f"  {i:>3}. {p.name}")

    # 1. Шрифт
    try:
        font_path = resolve_font_path()
        font = ImageFont.truetype(str(font_path), size=32)
        family, _ = font.getname()
        log(f"Шрифт: {family!r} ({font_path.name})")
    except Exception as exc:
        log(f"\nНе смог загрузить шрифт: {exc}")
        return 3

    # 2. Gemini OCR + порядок
    try:
        _ordered_imgs, ordered_texts = ai_pipeline(images, out_dir)
    except Exception as exc:
        log(f"\nОшибка Gemini: {exc}")
        return 2

    # 3. Параметры рендера
    width_cm, height_cm, font_size_pt = ask_render_params()

    # 4. Рендер текстов в PNG через TTF
    log()
    log("Рендерю текст в изображения через ваш шрифт...")
    rendered_images = render_texts_to_images(
        ordered_texts, font_path, out_dir,
        width_cm=width_cm,
        height_cm=height_cm,
        font_size_pt=font_size_pt,
    )

    # 5. Booklet docx из отрендеренных картинок
    generate_booklet(rendered_images, out_dir, width_cm, height_cm, layout=layout)

    log()
    log(f"Готово. Файлы здесь: {out_dir}")
    log(f"Отрендеренные страницы: {out_dir / 'rendered'}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
